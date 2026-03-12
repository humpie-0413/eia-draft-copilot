"""DOCX/PDF 출력 서비스.

초안 뼈대(DraftScaffold) 데이터를 기반으로 DOCX 및 PDF 문서를 생성한다.
- DOCX: python-docx 사용
- PDF: reportlab 사용 (한글 폰트 지원)
"""

from __future__ import annotations

import io
import os
import uuid
from datetime import datetime, timezone

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy.ext.asyncio import AsyncSession

from app.services.draft_scaffold import (
    DraftScaffold,
    ScaffoldSection,
    generate_draft_scaffold,
)
from app.services.qa_engine import run_qa


async def generate_docx(
    db: AsyncSession,
    project_id: uuid.UUID,
    project_name: str,
    *,
    skip_qa_check: bool = False,
) -> tuple[io.BytesIO, str]:
    """DOCX 문서를 생성하여 BytesIO와 파일명을 반환한다.

    Args:
        db: DB 세션
        project_id: 프로젝트 ID
        project_name: 프로젝트 이름 (표지 · 파일명에 사용)
        skip_qa_check: True면 QA 검사를 건너뜀 (테스트용)

    Returns:
        (BytesIO 버퍼, 파일명) 튜플

    Raises:
        ValueError: critical QA 이슈가 있어 export가 차단된 경우
    """
    # Export Gate: QA 검사
    if not skip_qa_check:
        qa_result = await run_qa(db, project_id)
        if not qa_result.export_ready:
            raise ValueError(
                f"critical 이슈 {qa_result.summary.critical_count}건이 "
                f"남아 있어 export가 차단되었습니다."
            )

    # 초안 뼈대 생성
    scaffold = await generate_draft_scaffold(db, project_id)

    # DOCX 문서 생성
    doc = _build_docx(scaffold, project_name)

    # BytesIO로 직렬화
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 파일명 생성 (HTTP 헤더 latin-1 호환을 위해 ASCII 안전 처리)
    timestamp = datetime.now(tz=timezone.utc).strftime("%Y%m%d_%H%M%S")
    # 한글 등 비ASCII 문자를 제거하고 ASCII 안전 파일명 생성
    safe_name = "".join(
        c if c.isascii() and c.isalnum() or c in "-_" else "_"
        for c in project_name.replace(" ", "_")
    )[:50].strip("_") or "draft"
    filename = f"EIA_{safe_name}_{timestamp}.docx"

    return buffer, filename


def _build_docx(scaffold: DraftScaffold, project_name: str) -> Document:
    """DraftScaffold로부터 python-docx Document를 생성한다."""
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ── 표지 ──
    _add_cover_page(doc, project_name, scaffold.generated_at)

    # ── 목차 페이지 ──
    _add_toc_page(doc, scaffold)

    # ── 섹션별 내용 ──
    for section in scaffold.sections:
        _add_section(doc, section)

    return doc


def _add_cover_page(doc: Document, project_name: str, generated_at: str) -> None:
    """표지 페이지를 추가한다."""
    # 빈 줄 추가 (상단 여백)
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("환경영향평가서")
    run.font.size = Pt(28)
    run.bold = True

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(project_name)
    run.font.size = Pt(18)

    doc.add_paragraph("")
    doc.add_paragraph("")

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = generated_at[:10] if len(generated_at) >= 10 else generated_at
    run = date_para.add_run(f"생성일: {date_str}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("EIA Draft Copilot에 의해 자동 생성됨")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()


def _add_toc_page(doc: Document, scaffold: DraftScaffold) -> None:
    """목차 페이지를 추가한다."""
    heading = doc.add_heading("목 차", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    for section in scaffold.sections:
        entry_count = len(section.evidence_entries)
        status = "근거 있음" if entry_count > 0 else "미수집"
        para = doc.add_paragraph(
            f"{section.order}. {section.title} — {status} ({entry_count}건)"
        )
        para.style = doc.styles["List Number"]

    doc.add_page_break()


def _add_section(doc: Document, section: ScaffoldSection) -> None:
    """개별 섹션을 DOCX에 추가한다."""
    doc.add_heading(f"{section.order}. {section.title}", level=1)
    doc.add_paragraph(section.description).italic = True

    if not section.evidence_entries:
        para = doc.add_paragraph(
            f"{section.title} 분야에 대한 수집된 증거 데이터가 없습니다."
        )
        para.runs[0].font.color.rgb = RGBColor(180, 0, 0)
        doc.add_paragraph("")
        return

    # 요약문 추가
    doc.add_heading("현황 요약", level=2)
    for line in section.summary_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)

    # 근거 데이터 테이블
    doc.add_heading("근거 데이터 목록", level=2)
    _add_evidence_table(doc, section)

    doc.add_paragraph("")


def _add_evidence_table(doc: Document, section: ScaffoldSection) -> None:
    """근거 데이터를 테이블로 추가한다."""
    entries = section.evidence_entries
    if not entries:
        return

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 헤더
    headers = ["지표", "측정값", "단위", "관측일"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)

    # 데이터 행
    for entry in entries:
        row = table.add_row()
        row.cells[0].text = entry.indicator
        row.cells[1].text = entry.value
        row.cells[2].text = entry.unit or "-"
        row.cells[3].text = entry.observed_at[:10] if entry.observed_at else "-"

        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

    # 컬럼 너비 설정
    widths = [Cm(4), Cm(5), Cm(3), Cm(3)]
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width


# ═══════════════════════════════════════════════════════════════
# PDF 출력 (reportlab)
# ═══════════════════════════════════════════════════════════════

def _register_korean_font() -> str:
    """한글 폰트를 등록하고 폰트 이름을 반환한다.

    시스템에 맑은 고딕(malgun.ttf) 또는 나눔고딕(NanumGothic.ttf)이
    있으면 TTFont로 등록하고, 없으면 CID 한글 폰트를 사용한다.
    """
    # 이미 등록된 경우 재등록 방지
    registered = pdfmetrics.getRegisteredFontNames()

    # 1순위: 맑은 고딕
    for font_path in [
        "C:/Windows/Fonts/malgun.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/nanum/NanumGothic.ttf",
    ]:
        if os.path.exists(font_path):
            font_name = "MalgunGothic" if "malgun" in font_path else "NanumGothic"
            if font_name not in registered:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
            return font_name

    # 2순위: CID 한글 폰트 (reportlab 내장)
    cid_name = "HYGothic-Medium"
    if cid_name not in registered:
        pdfmetrics.registerFont(UnicodeCIDFont(cid_name))
    return cid_name


def _get_pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    """PDF용 스타일 사전을 생성한다."""
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "KoTitle",
            parent=base["Title"],
            fontName=font_name,
            fontSize=28,
            leading=34,
            alignment=1,  # CENTER
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "KoSubtitle",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            alignment=1,
            spaceAfter=8,
        ),
        "heading1": ParagraphStyle(
            "KoH1",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=16,
            leading=22,
            spaceBefore=20,
            spaceAfter=10,
        ),
        "heading2": ParagraphStyle(
            "KoH2",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=13,
            leading=18,
            spaceBefore=14,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "KoBody",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=15,
            spaceAfter=4,
        ),
        "small": ParagraphStyle(
            "KoSmall",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=8,
            leading=12,
        ),
        "center": ParagraphStyle(
            "KoCenter",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=15,
            alignment=1,
        ),
        "center_gray": ParagraphStyle(
            "KoCenterGray",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=15,
            alignment=1,
            textColor=colors.Color(0.6, 0.6, 0.6),
        ),
        "italic_desc": ParagraphStyle(
            "KoItalicDesc",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            textColor=colors.Color(0.3, 0.3, 0.3),
            spaceAfter=8,
        ),
        "no_data": ParagraphStyle(
            "KoNoData",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            textColor=colors.Color(0.7, 0, 0),
        ),
    }


async def generate_pdf(
    db: AsyncSession,
    project_id: uuid.UUID,
    project_name: str,
    *,
    skip_qa_check: bool = False,
) -> tuple[io.BytesIO, str]:
    """PDF 문서를 생성하여 BytesIO와 파일명을 반환한다.

    내용 구조는 DOCX와 동일하다 (표지 → 목차 → 섹션별 요약문 + 근거 테이블).

    Args:
        db: DB 세션
        project_id: 프로젝트 ID
        project_name: 프로젝트 이름
        skip_qa_check: True면 QA 검사 건너뜀 (테스트용)

    Returns:
        (BytesIO 버퍼, 파일명) 튜플

    Raises:
        ValueError: critical QA 이슈로 export 차단 시
    """
    # Export Gate: QA 검사 (DOCX와 동일 로직)
    if not skip_qa_check:
        qa_result = await run_qa(db, project_id)
        if not qa_result.export_ready:
            raise ValueError(
                f"critical 이슈 {qa_result.summary.critical_count}건이 "
                f"남아 있어 export가 차단되었습니다."
            )

    # 초안 뼈대 생성
    scaffold = await generate_draft_scaffold(db, project_id)

    # PDF 문서 생성
    buffer = _build_pdf(scaffold, project_name)

    # 파일명 생성
    timestamp = datetime.now(tz=timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(
        c if c.isascii() and c.isalnum() or c in "-_" else "_"
        for c in project_name.replace(" ", "_")
    )[:50].strip("_") or "draft"
    filename = f"EIA_{safe_name}_{timestamp}.pdf"

    return buffer, filename


def _build_pdf(scaffold: DraftScaffold, project_name: str) -> io.BytesIO:
    """DraftScaffold로부터 reportlab PDF 문서를 생성한다."""
    font_name = _register_korean_font()
    styles = _get_pdf_styles(font_name)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        title=f"환경영향평가서 — {project_name}",
        author="EIA Draft Copilot",
    )

    story: list = []

    # ── 표지 ──
    _pdf_add_cover(story, styles, project_name, scaffold.generated_at)

    # ── 목차 ──
    _pdf_add_toc(story, styles, scaffold)

    # ── 섹션별 내용 ──
    for section in scaffold.sections:
        _pdf_add_section(story, styles, font_name, section)

    doc.build(story)
    buffer.seek(0)
    return buffer


def _pdf_add_cover(
    story: list,
    styles: dict[str, ParagraphStyle],
    project_name: str,
    generated_at: str,
) -> None:
    """표지 페이지를 추가한다."""
    story.append(Spacer(1, 6 * cm))
    story.append(Paragraph("환경영향평가서", styles["title"]))
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(project_name, styles["subtitle"]))
    story.append(Spacer(1, 2 * cm))

    date_str = generated_at[:10] if len(generated_at) >= 10 else generated_at
    story.append(Paragraph(f"생성일: {date_str}", styles["center"]))
    story.append(Spacer(1, 0.5 * cm))
    story.append(
        Paragraph("EIA Draft Copilot에 의해 자동 생성됨", styles["center_gray"])
    )
    story.append(PageBreak())


def _pdf_add_toc(
    story: list,
    styles: dict[str, ParagraphStyle],
    scaffold: DraftScaffold,
) -> None:
    """목차 페이지를 추가한다."""
    story.append(Paragraph("목 차", styles["heading1"]))
    story.append(Spacer(1, 0.5 * cm))

    for section in scaffold.sections:
        entry_count = len(section.evidence_entries)
        status = "근거 있음" if entry_count > 0 else "미수집"
        story.append(
            Paragraph(
                f"{section.order}. {section.title} — {status} ({entry_count}건)",
                styles["body"],
            )
        )

    story.append(PageBreak())


def _pdf_add_section(
    story: list,
    styles: dict[str, ParagraphStyle],
    font_name: str,
    section: ScaffoldSection,
) -> None:
    """개별 섹션을 PDF에 추가한다."""
    story.append(
        Paragraph(f"{section.order}. {section.title}", styles["heading1"])
    )
    story.append(Paragraph(section.description, styles["italic_desc"]))

    if not section.evidence_entries:
        story.append(
            Paragraph(
                f"{section.title} 분야에 대한 수집된 증거 데이터가 없습니다.",
                styles["no_data"],
            )
        )
        story.append(Spacer(1, 0.5 * cm))
        return

    # 요약문
    story.append(Paragraph("현황 요약", styles["heading2"]))
    for line in section.summary_text.split("\n"):
        if line.strip():
            story.append(Paragraph(line, styles["body"]))

    # 근거 데이터 테이블
    story.append(Paragraph("근거 데이터 목록", styles["heading2"]))
    _pdf_add_evidence_table(story, font_name, section)
    story.append(Spacer(1, 0.5 * cm))


def _pdf_add_evidence_table(
    story: list,
    font_name: str,
    section: ScaffoldSection,
) -> None:
    """근거 데이터를 테이블로 추가한다."""
    entries = section.evidence_entries
    if not entries:
        return

    header_style = ParagraphStyle(
        "TableHeader",
        fontName=font_name,
        fontSize=9,
        leading=12,
        alignment=1,
        textColor=colors.white,
    )
    cell_style = ParagraphStyle(
        "TableCell",
        fontName=font_name,
        fontSize=8,
        leading=11,
    )

    # 테이블 데이터 구성
    data = [
        [
            Paragraph("지표", header_style),
            Paragraph("측정값", header_style),
            Paragraph("단위", header_style),
            Paragraph("관측일", header_style),
        ]
    ]
    for entry in entries:
        observed = entry.observed_at[:10] if entry.observed_at else "-"
        data.append([
            Paragraph(entry.indicator, cell_style),
            Paragraph(entry.value, cell_style),
            Paragraph(entry.unit or "-", cell_style),
            Paragraph(observed, cell_style),
        ])

    # 컬럼 너비 (A4 기준, 여백 제외 약 17cm)
    col_widths = [5 * cm, 5 * cm, 3.5 * cm, 3.5 * cm]
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        # 헤더 배경
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.2, 0.3, 0.5)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        # 격자선
        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0.7, 0.7, 0.7)),
        # 패딩
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        # 줄무늬 배경
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.Color(0.95, 0.95, 0.97)]),
        # 정렬
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))

    story.append(table)

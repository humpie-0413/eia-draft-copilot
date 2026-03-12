"""전체 워크플로우 End-to-End 통합 테스트.

전체 흐름을 하나의 시나리오로 관통한다:
1. 프로젝트 생성 (GeoJSON geometry 포함)
2. geometry 검증 및 저장 확인
3. 증거(evidence) 수동 추가 (대기, 수질, 소음진동 등 최소 3개 섹션)
4. screening_only 필터 동작 확인
5. 유사사례 매칭 실행 및 결과 확인
6. 섹션 플래너 충족도 계산 확인
7. 초안 뼈대 생성 확인 (evidence 기반 근거 텍스트 배치)
8. QA 규칙 실행 및 이슈 목록 확인
9. critical 이슈 있을 때 export 차단 확인
10. 이슈 해소 후 DOCX export 성공 확인

실행 방법:
  cd backend
  pytest tests/test_e2e.py -v --tb=short

사전 조건:
  - PostgreSQL + PostGIS 실행 중 (localhost:5432)
  - eia_copilot_test DB 생성 + PostGIS 확장 활성화:
    CREATE DATABASE eia_copilot_test;
    \\c eia_copilot_test
    CREATE EXTENSION IF NOT EXISTS postgis;
"""

import io
from typing import Any

import pytest
from httpx import AsyncClient


# ── 테스트 데이터 ──

SAMPLE_POLYGON = {
    "type": "Polygon",
    "coordinates": [
        [
            [127.0, 37.5],
            [127.1, 37.5],
            [127.1, 37.6],
            [127.0, 37.6],
            [127.0, 37.5],
        ]
    ],
}

# 대기질 증거 — 6개 필수 지표 전부 충족
AIR_EVIDENCES = [
    {"category": "air_quality", "indicator": "PM10_연평균", "value": "45", "numeric_value": 45.0, "unit": "ug/m3"},
    {"category": "air_quality", "indicator": "PM2.5_연평균", "value": "23", "numeric_value": 23.0, "unit": "ug/m3"},
    {"category": "air_quality", "indicator": "NO2_연평균", "value": "0.028", "numeric_value": 0.028, "unit": "ppm"},
    {"category": "air_quality", "indicator": "SO2_연평균", "value": "0.004", "numeric_value": 0.004, "unit": "ppm"},
    {"category": "air_quality", "indicator": "CO_연평균", "value": "0.5", "numeric_value": 0.5, "unit": "ppm"},
    {"category": "air_quality", "indicator": "O3_연평균", "value": "0.032", "numeric_value": 0.032, "unit": "ppm"},
]

# 스크리닝 전용 대기질 증거
AIR_SCREENING = [
    {"category": "air_quality", "indicator": "PM10_일평균", "value": "55", "numeric_value": 55.0, "unit": "ug/m3", "screening_only": True},
]

# 수질 증거 — 6개 필수 지표 전부 충족
WATER_EVIDENCES = [
    {"category": "water_quality", "indicator": "BOD", "value": "2.1", "numeric_value": 2.1, "unit": "mg/L"},
    {"category": "water_quality", "indicator": "COD", "value": "4.3", "numeric_value": 4.3, "unit": "mg/L"},
    {"category": "water_quality", "indicator": "SS", "value": "12", "numeric_value": 12.0, "unit": "mg/L"},
    {"category": "water_quality", "indicator": "T-N", "value": "1.8", "numeric_value": 1.8, "unit": "mg/L"},
    {"category": "water_quality", "indicator": "T-P", "value": "0.05", "numeric_value": 0.05, "unit": "mg/L"},
    {"category": "water_quality", "indicator": "DO", "value": "8.5", "numeric_value": 8.5, "unit": "mg/L"},
]

# 소음진동 증거 — 2/3 충족 (진동_Lv_주간 누락)
NOISE_EVIDENCES_PARTIAL = [
    {"category": "noise_vibration", "indicator": "소음_Leq_주간", "value": "58", "numeric_value": 58.0, "unit": "dB(A)"},
    {"category": "noise_vibration", "indicator": "소음_Leq_야간", "value": "49", "numeric_value": 49.0, "unit": "dB(A)"},
]

# 생태 증거 — 5개 필수 지표 전부 충족
ECOLOGY_EVIDENCES = [
    {"category": "ecology", "indicator": "식물상_종수", "value": "245", "numeric_value": 245.0, "unit": "종"},
    {"category": "ecology", "indicator": "동물상_종수", "value": "78", "numeric_value": 78.0, "unit": "종"},
    {"category": "ecology", "indicator": "법정보호종", "value": "2", "numeric_value": 2.0, "unit": "종"},
    {"category": "ecology", "indicator": "비오톱_유형", "value": "자연림"},
    {"category": "ecology", "indicator": "녹지자연도", "value": "7등급"},
]

# 소음진동 누락 지표 보충
NOISE_MISSING = [
    {"category": "noise_vibration", "indicator": "진동_Lv_주간", "value": "63", "numeric_value": 63.0, "unit": "dB(V)"},
]

SIMILAR_CASE_DATA = {
    "name": "E2E 유사사례",
    "description": "산업단지 개발 유사사례",
    "project_type": "industrial",
    "location": {"type": "Point", "coordinates": [127.05, 37.55]},
    "area_sqm": 500000.0,
    "summary": "산업단지 조성에 따른 환경영향평가 사례",
    "key_findings": {"f1": "대기질 영향 미미", "f2": "수질 관리 필요"},
    "evidence_categories": ["air_quality", "water_quality", "noise_vibration"],
}


# ── 헬퍼 함수 ──

async def _create_evidence(client: AsyncClient, project_id: str, ev: dict[str, Any]) -> dict:
    """증거 데이터 생성 헬퍼."""
    payload = {"project_id": project_id, "screening_only": False, **ev}
    resp = await client.post("/api/v1/evidences", json=payload)
    assert resp.status_code == 201, f"증거 생성 실패: {ev['indicator']} -> {resp.text}"
    return resp.json()


# ── 테스트 ──

@pytest.mark.asyncio
async def test_e2e_full_workflow(client: AsyncClient):
    """전체 MVP 워크플로우를 하나의 시나리오로 관통하는 E2E 테스트."""

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 1: 프로젝트 생성 (GeoJSON geometry 포함)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.post(
        "/api/v1/projects",
        json={
            "name": "E2E 통합 테스트 프로젝트",
            "description": "실사용 검증용",
            "project_type": "industrial",
            "geometry": SAMPLE_POLYGON,
        },
    )
    assert resp.status_code == 201
    project = resp.json()
    project_id = project["id"]
    assert project["status"] == "draft"
    assert project["project_type"] == "industrial"

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 2: geometry 검증 및 저장 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.get(f"/api/v1/projects/{project_id}")
    assert resp.status_code == 200
    project_detail = resp.json()
    assert project_detail["geometry"]["type"] == "Polygon"
    assert len(project_detail["geometry"]["coordinates"][0]) == 5

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 3: 증거(evidence) 수동 추가 — 대기, 수질, 소음진동 (3개 섹션)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    for ev in AIR_EVIDENCES:
        await _create_evidence(client, project_id, ev)
    for ev in AIR_SCREENING:
        await _create_evidence(client, project_id, ev)
    for ev in WATER_EVIDENCES:
        await _create_evidence(client, project_id, ev)
    for ev in NOISE_EVIDENCES_PARTIAL:
        await _create_evidence(client, project_id, ev)

    # 총 15건 = 대기 6 + 스크리닝 1 + 수질 6 + 소음 2
    resp = await client.get(f"/api/v1/evidences?project_id={project_id}")
    assert resp.status_code == 200
    assert resp.json()["total"] == 15

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 4: screening_only 필터 동작 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.get(f"/api/v1/evidences?project_id={project_id}&screening_only=true")
    assert resp.status_code == 200
    assert resp.json()["total"] == 1, "스크리닝 전용 증거는 1건이어야 함"

    resp = await client.get(f"/api/v1/evidences?project_id={project_id}&screening_only=false")
    assert resp.status_code == 200
    assert resp.json()["total"] == 14, "본 평가 증거는 14건이어야 함"

    resp = await client.get(f"/api/v1/evidences?project_id={project_id}&category=air_quality")
    assert resp.status_code == 200
    assert resp.json()["total"] == 7, "대기질 전체(스크리닝 포함) 7건"

    resp = await client.get(
        f"/api/v1/evidences?project_id={project_id}&category=air_quality&screening_only=false"
    )
    assert resp.status_code == 200
    assert resp.json()["total"] == 6, "대기질 본 평가만 6건"

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 5: 유사사례 매칭 실행 및 결과 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 유사사례 등록
    resp = await client.post("/api/v1/similar-cases", json=SIMILAR_CASE_DATA)
    assert resp.status_code == 201
    case = resp.json()
    assert case["project_type"] == "industrial"

    # 매칭 실행
    resp = await client.get(f"/api/v1/similar-cases/match/{project_id}")
    assert resp.status_code == 200
    matches = resp.json()
    assert matches["total"] >= 1, "최소 1건의 매칭 결과가 있어야 함"
    best = matches["matches"][0]
    assert best["type_score"] == 1.0, "같은 사업 유형이므로 type_score=1.0"
    assert best["overall_score"] > 0.5, "종합 유사도 0.5 이상"
    assert best["similar_case"]["name"] == "E2E 유사사례"

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 6: 섹션 플래너 충족도 계산 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.get(f"/api/v1/projects/{project_id}/sections/status")
    assert resp.status_code == 200
    sections = resp.json()
    section_map = {s["section_key"]: s for s in sections["sections"]}

    # 대기질: 6/6 충족 → complete
    assert section_map["air_quality"]["status"] == "complete"
    assert section_map["air_quality"]["coverage_ratio"] == 1.0
    assert section_map["air_quality"]["fulfilled_count"] == 6

    # 수질: 6/6 충족 → complete
    assert section_map["water_quality"]["status"] == "complete"
    assert section_map["water_quality"]["coverage_ratio"] == 1.0

    # 소음진동: 2/3 충족 → partial
    assert section_map["noise_vibration"]["status"] == "partial"
    assert section_map["noise_vibration"]["coverage_ratio"] < 1.0
    assert section_map["noise_vibration"]["fulfilled_count"] == 2

    # 생태: 0건 → empty
    assert section_map["ecology"]["status"] == "empty"
    assert section_map["ecology"]["total_evidence_count"] == 0

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 7: 초안 뼈대 생성 확인 (evidence 기반 근거 텍스트 배치)
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.get(f"/api/v1/projects/{project_id}/sections/scaffold")
    assert resp.status_code == 200
    scaffold = resp.json()
    assert scaffold["total_evidence_count"] == 14, "스크리닝 제외 본 평가 14건"

    scaffold_map = {s["section_key"]: s for s in scaffold["sections"]}

    # 대기질 뼈대: 6건 근거, 요약 텍스트에 PM10 언급
    assert len(scaffold_map["air_quality"]["evidence_entries"]) == 6
    assert "PM10" in scaffold_map["air_quality"]["summary_text"]

    # 수질 뼈대: 6건 근거
    assert len(scaffold_map["water_quality"]["evidence_entries"]) == 6

    # 생태 뼈대: 0건 (미수집)
    assert len(scaffold_map["ecology"]["evidence_entries"]) == 0
    assert "없습니다" in scaffold_map["ecology"]["summary_text"]

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 8: QA 규칙 실행 및 이슈 목록 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.get(f"/api/v1/projects/{project_id}/qa")
    assert resp.status_code == 200
    qa = resp.json()

    assert qa["export_ready"] is False, "critical 이슈가 있으므로 export 불가"
    assert qa["summary"]["critical_count"] >= 2, "생태 R001 + 소음 R002 최소 2건"

    # R001: 생태 섹션 증거 없음 (critical)
    ecology_r001 = [
        i for i in qa["issues"]
        if i["rule_id"] == "R001" and i["section_key"] == "ecology"
    ]
    assert len(ecology_r001) == 1
    assert ecology_r001[0]["severity"] == "critical"

    # R002: 소음진동 필수 지표 누락 (critical)
    noise_r002 = [
        i for i in qa["issues"]
        if i["rule_id"] == "R002" and i["section_key"] == "noise_vibration"
    ]
    assert len(noise_r002) == 1
    assert noise_r002[0]["severity"] == "critical"
    assert "진동_Lv_주간" in noise_r002[0]["indicators"]

    # export-ready 간략 확인 API
    resp = await client.get(f"/api/v1/projects/{project_id}/qa/export-ready")
    assert resp.status_code == 200
    assert resp.json()["export_ready"] is False

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 9: critical 이슈 있을 때 export 차단 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.post(f"/api/v1/projects/{project_id}/export/docx")
    assert resp.status_code == 422, "critical 이슈가 있으면 export가 차단(422)되어야 함"
    assert "critical" in resp.json()["detail"].lower()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 10: 이슈 해소 후 DOCX export 성공 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 생태 증거 추가 (R001 해소)
    for ev in ECOLOGY_EVIDENCES:
        await _create_evidence(client, project_id, ev)

    # 소음진동 누락 지표 보충 (R002 해소)
    for ev in NOISE_MISSING:
        await _create_evidence(client, project_id, ev)

    # QA 재검사: critical 이슈 0건
    resp = await client.get(f"/api/v1/projects/{project_id}/qa")
    assert resp.status_code == 200
    qa2 = resp.json()
    assert qa2["summary"]["critical_count"] == 0, "critical 이슈가 모두 해소되어야 함"
    assert qa2["export_ready"] is True

    # DOCX export 성공
    resp = await client.post(f"/api/v1/projects/{project_id}/export/docx")
    assert resp.status_code == 200, f"DOCX export 실패: {resp.text}"
    assert "wordprocessingml.document" in resp.headers["content-type"]
    assert "attachment" in resp.headers["content-disposition"]

    # DOCX 내용 검증 (python-docx)
    from docx import Document

    doc = Document(io.BytesIO(resp.content))
    all_text = " ".join(p.text for p in doc.paragraphs)

    assert "환경영향평가서" in all_text, "표지에 '환경영향평가서' 제목 필요"
    assert "대기질" in all_text, "대기질 섹션이 DOCX에 포함되어야 함"
    assert "수질" in all_text, "수질 섹션이 DOCX에 포함되어야 함"
    assert "생태" in all_text, "생태 섹션이 DOCX에 포함되어야 함"
    assert len(doc.tables) >= 3, "근거 데이터 테이블이 최소 3개 필요 (대기, 수질, 소음)"

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 단계 11: PDF export 성공 확인
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    resp = await client.get(f"/api/v1/projects/{project_id}/export/pdf")
    assert resp.status_code == 200, f"PDF export 실패: {resp.text}"
    assert resp.headers["content-type"] == "application/pdf"
    assert "attachment" in resp.headers["content-disposition"]
    assert resp.content[:5] == b"%PDF-", "유효한 PDF 파일이어야 함"
    assert len(resp.content) > 5000, "충분한 내용이 담긴 PDF여야 함"
